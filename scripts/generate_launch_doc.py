from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Estilos generales ──────────────────────────────────────────────
def set_font(run, size=11, bold=False, color=None, italic=False):
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p

def add_colored_paragraph(doc, text, color=(0,0,0), size=11, bold=False, space_before=0, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, color=color)
    return p

def shade_cell(cell, fill_hex):
    """Set background color of a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def add_status_table(doc, rows_data, headers):
    """rows_data: list of (col1, col2, color_hex_for_col2)"""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        shade_cell(hdr[i], '2C3E50')
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.size = Pt(10)
        if not hdr[i].paragraphs[0].runs:
            r = hdr[i].paragraphs[0].add_run(h)
            r.bold = True
            r.font.color.rgb = RGBColor(255, 255, 255)
            r.font.size = Pt(10)

    for row_data in rows_data:
        row = table.add_row().cells
        for i, cell_data in enumerate(row_data[:-1]):
            row[i].text = str(cell_data)
            for para in row[i].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
        # Last column with color
        last_idx = len(row_data) - 2
        status_text = row_data[-2] if len(row_data) > 1 else ''
        fill = row_data[-1]
        last_cell = row[len(row_data) - 2]
        last_cell.text = status_text
        shade_cell(last_cell, fill)
        for para in last_cell.paragraphs:
            for run in para.runs:
                run.font.size = Pt(10)
                run.bold = True

    return table

# ── PORTADA ────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('COUTAZO')
run.bold = True
run.font.size = Pt(36)
run.font.color.rgb = RGBColor(0, 102, 204)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run('Plan de Lanzamiento a Producción')
run2.font.size = Pt(20)
run2.font.color.rgb = RGBColor(44, 62, 80)

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = p3.add_run(f'Fecha: {datetime.date.today().strftime("%d de %B de %Y")}')
run3.font.size = Pt(12)
run3.font.color.rgb = RGBColor(127, 140, 141)

doc.add_paragraph()

# Estado general
box = doc.add_paragraph()
box.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = box.add_run('⚠  ESTADO GENERAL: NO LISTO PARA PRODUCCIÓN')
r.bold = True
r.font.size = Pt(13)
r.font.color.rgb = RGBColor(192, 57, 43)

doc.add_page_break()

# ── ÍNDICE ─────────────────────────────────────────────────────────
add_heading(doc, '📋 Contenido', level=1)
items_indice = [
    '1. Diagnóstico General',
    '2. Bloqueos Críticos — Variables de Entorno',
    '3. Bloqueos Críticos — Pagos (Stripe y PayU)',
    '4. Bloqueos Críticos — Edge Functions de Supabase',
    '5. Implementaciones Faltantes — Cron Jobs',
    '6. Implementaciones Faltantes — Páginas y Rutas',
    '7. Lo que SÍ está listo',
    '8. Plan de Acción por Semanas',
    '9. Estimación de Tiempo',
    '10. Checklist Final Pre-Lanzamiento',
]
for item in items_indice:
    p = doc.add_paragraph(item, style='List Number')
    p.paragraph_format.space_after = Pt(3)

doc.add_page_break()

# ── SECCIÓN 1 — DIAGNÓSTICO ────────────────────────────────────────
add_heading(doc, '1. Diagnóstico General', level=1)
add_colored_paragraph(doc,
    'Tras una auditoría completa del código fuente, base de datos, integraciones de pago, '
    'cron jobs y configuración de entorno, se identificaron los siguientes hallazgos:',
    size=11)

doc.add_paragraph()

# Tabla resumen
add_heading(doc, 'Resumen por Área', level=2)
summary_rows = [
    ('Schema de base de datos (17 tablas + RLS)', '✓ Completo',     'C8E6C9', 'C8E6C9'),
    ('Sistema de autenticación',                  '✓ Completo',     'C8E6C9', 'C8E6C9'),
    ('Middleware (rate limiting, bot detection)',  '✓ Completo',     'C8E6C9', 'C8E6C9'),
    ('Páginas de marketing / home',               '✓ Completo',     'C8E6C9', 'C8E6C9'),
    ('Dashboard de usuario',                      '✓ Completo',     'C8E6C9', 'C8E6C9'),
    ('Precios / premium gate',                    '✓ Completo',     'C8E6C9', 'C8E6C9'),
    ('Leaderboard',                               '✓ Completo',     'C8E6C9', 'C8E6C9'),
    ('Lógica matemática (value bets, Kelly, Poisson)', '✓ Completo','C8E6C9', 'C8E6C9'),
    ('Backtesting (Pro)',                         '✓ Completo',     'C8E6C9', 'C8E6C9'),
    ('API Key (Pro)',                             '✓ Completo',     'C8E6C9', 'C8E6C9'),
    ('Integración Stripe',                        '⚠ Incompleto',   'FFF9C4', 'FFF9C4'),
    ('Integración PayU',                          '⚠ Incompleto',   'FFF9C4', 'FFF9C4'),
    ('Cron Jobs (5 de 11 implementados)',          '⚠ Parcial',      'FFF9C4', 'FFF9C4'),
    ('Edge Functions de Supabase',                '✗ Sin desplegar','FFCDD2', 'FFCDD2'),
    ('Variables de entorno',                      '✗ 10+ sin llenar','FFCDD2','FFCDD2'),
    ('Páginas Pro (telegram, webhooks, mis-picks)','✗ Incompleto',  'FFCDD2', 'FFCDD2'),
    ('Rastreo de afiliados (ruta API)',            '✗ Falta ruta',   'FFCDD2', 'FFCDD2'),
    ('Tests E2E',                                 '✗ Sin ejecutar', 'FFCDD2', 'FFCDD2'),
]

table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
hdr = table.rows[0].cells
for i, h in enumerate(['Área', 'Estado']):
    shade_cell(hdr[i], '2C3E50')
    r = hdr[i].paragraphs[0].add_run(h)
    r.bold = True
    r.font.color.rgb = RGBColor(255, 255, 255)
    r.font.size = Pt(10)

for area, estado, fill, _ in summary_rows:
    row = table.add_row().cells
    row[0].text = area
    row[1].text = estado
    shade_cell(row[1], fill)
    for para in row[0].paragraphs:
        for run in para.runs:
            run.font.size = Pt(10)
    for para in row[1].paragraphs:
        for run in para.runs:
            run.font.size = Pt(10)
            run.bold = True

doc.add_page_break()

# ── SECCIÓN 2 — VARIABLES DE ENTORNO ──────────────────────────────
add_heading(doc, '2. Bloqueos Críticos — Variables de Entorno', level=1)
add_colored_paragraph(doc,
    'Las siguientes variables en .env.local contienen valores placeholder o no existen. '
    'Sin ellas el sistema no puede procesar pagos, ejecutar crons ni enviar notificaciones.',
    size=11)
doc.add_paragraph()

env_sections = [
    ('Stripe', [
        ('STRIPE_SECRET_KEY', 'sk_test_... (placeholder)', 'Crear cuenta Stripe → Developers → API Keys'),
        ('NEXT_PUBLIC_STRIPE_PUBLISHABLE_KEY', 'pk_test_... (placeholder)', 'Mismo panel de Stripe'),
        ('STRIPE_WEBHOOK_SECRET', 'whsec_... (placeholder)', 'Stripe → Webhooks → Signing secret'),
        ('STRIPE_PRICE_ID_MONTHLY', 'price_... (placeholder)', 'Crear producto en Stripe → copiar Price ID'),
        ('STRIPE_PRICE_ID_YEARLY', 'price_... (placeholder)', 'Crear producto en Stripe → copiar Price ID'),
        ('STRIPE_PRICE_ID_PRO_MONTHLY', 'NO EXISTE en .env', 'Crear producto Pro mensual en Stripe'),
        ('STRIPE_PRICE_ID_PRO_YEARLY', 'NO EXISTE en .env', 'Crear producto Pro anual en Stripe'),
    ]),
    ('Seguridad / Crons', [
        ('CRON_SECRET', 'Contiene código shell sin ejecutar', 'Ejecutar: openssl rand -hex 32'),
    ]),
    ('PayU', [
        ('PAYU_RESPONSE_URL', 'No existe en .env.local', 'https://[tu-dominio.com]/pago/respuesta'),
        ('PAYU_NOTIFICATION_URL', 'No existe en .env.local', 'https://[tu-dominio.com]/api/webhooks/payu'),
    ]),
    ('APIs Externas', [
        ('ODDS_API_KEY', 'Placeholder (xxxxxxx)', 'Registrarse en the-odds-api.com'),
        ('NEWS_API_KEY', 'Placeholder (xxxxxxx)', 'Registrarse en newsapi.org (plan gratuito disponible)'),
    ]),
    ('Notificaciones', [
        ('NEXT_PUBLIC_ONESIGNAL_APP_ID', 'Placeholder (...)', 'Crear app en onesignal.com'),
        ('ONESIGNAL_REST_API_KEY', 'Placeholder (...)', 'Panel OneSignal → Settings → Keys & IDs'),
        ('TELEGRAM_BOT_TOKEN', 'NO EXISTE en .env', 'Crear bot con @BotFather en Telegram'),
    ]),
    ('Analytics', [
        ('NEXT_PUBLIC_POSTHOG_KEY', 'phc_... (placeholder)', 'Registrarse en posthog.com'),
        ('NEXT_PUBLIC_POSTHOG_HOST', 'Placeholder genérico', 'https://app.posthog.com o self-hosted'),
    ]),
]

for section_title, vars_list in env_sections:
    add_heading(doc, f'2.{env_sections.index((section_title, vars_list))+1}. {section_title}', level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, h in enumerate(['Variable', 'Estado actual', 'Cómo resolverlo']):
        shade_cell(hdr[i], '34495E')
        r = hdr[i].paragraphs[0].add_run(h)
        r.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)
        r.font.size = Pt(9)
    for var, estado, accion in vars_list:
        row = table.add_row().cells
        row[0].text = var
        row[1].text = estado
        row[2].text = accion
        shade_cell(row[1], 'FFCDD2')
        for j in range(3):
            for para in row[j].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)
    doc.add_paragraph()

doc.add_page_break()

# ── SECCIÓN 3 — PAGOS ─────────────────────────────────────────────
add_heading(doc, '3. Bloqueos Críticos — Pagos (Stripe y PayU)', level=1)

add_heading(doc, '3.1. Stripe', level=2)
add_colored_paragraph(doc, 'El código del webhook y checkout está implementado correctamente. Faltan los credenciales reales.', size=11)
doc.add_paragraph()

stripe_steps = [
    'Crear cuenta en stripe.com (o acceder a la existente).',
    'En el Dashboard de Stripe → Developers → API Keys: copiar la Publishable Key y la Secret Key.',
    'Ir a Products → Create product:',
    '   a. Producto "Premium Mensual" → precio $4.99 USD / mes → copiar el Price ID.',
    '   b. Producto "Premium Anual" → precio $49.99 USD / año → copiar el Price ID.',
    '   c. Producto "Pro Mensual" → precio $14.99 USD / mes → copiar el Price ID.',
    '   d. Producto "Pro Anual" → precio $149.99 USD / año → copiar el Price ID.',
    'En Developers → Webhooks → Add endpoint:',
    '   URL: https://[tu-dominio]/api/webhooks/stripe',
    '   Eventos a escuchar: checkout.session.completed, customer.subscription.created, customer.subscription.updated, customer.subscription.deleted, invoice.payment_failed',
    'Copiar el Signing Secret del webhook → llenar STRIPE_WEBHOOK_SECRET.',
    'Llenar todas las variables en .env.local y en Vercel → Settings → Environment Variables.',
    'Probar el flujo completo con una tarjeta de prueba: 4242 4242 4242 4242.',
]
for i, step in enumerate(stripe_steps, 1):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.space_after = Pt(3)
    if step.startswith('   '):
        p.paragraph_format.left_indent = Inches(0.5)
        p.add_run(step.strip()).font.size = Pt(10)
    else:
        p.add_run(step).font.size = Pt(10)

doc.add_paragraph()
add_heading(doc, '3.2. PayU', level=2)
add_colored_paragraph(doc, 'Credenciales sandbox configuradas. Para producción:', size=11)
doc.add_paragraph()
payu_steps = [
    'Registrar el dominio de producción en el panel de PayU → Merchant Center.',
    'Llenar PAYU_RESPONSE_URL = https://[tu-dominio]/pago/respuesta',
    'Llenar PAYU_NOTIFICATION_URL = https://[tu-dominio]/api/webhooks/payu',
    'Cambiar PAYU_ENV de "sandbox" a "production" al momento del lanzamiento real.',
    'Registrar la URL del webhook IPN en el panel de PayU (Merchant Center → Configuración → IPN).',
    'Probar con tarjeta de prueba de PayU antes de cambiar a producción.',
]
for i, step in enumerate(payu_steps, 1):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.space_after = Pt(3)
    p.add_run(step).font.size = Pt(10)

doc.add_page_break()

# ── SECCIÓN 4 — EDGE FUNCTIONS ────────────────────────────────────
add_heading(doc, '4. Bloqueos Críticos — Edge Functions de Supabase', level=1)
add_colored_paragraph(doc,
    'Las Edge Functions detect-value-bets y generate-parlays tienen el código escrito pero NO están '
    'desplegadas en Supabase. Vercel solo despliega el frontend Next.js; las funciones de Supabase '
    'deben desplegarse manualmente con el CLI.',
    size=11, color=(192, 57, 43))
doc.add_paragraph()

ef_steps = [
    'Instalar el CLI de Supabase: npm install -g supabase',
    'Autenticarse: supabase login',
    'Vincular el proyecto: supabase link --project-ref [PROJECT_ID]',
    'Desplegar la función detect-value-bets: supabase functions deploy detect-value-bets',
    'Desplegar la función generate-parlays: supabase functions deploy generate-parlays',
    'Verificar en Supabase Dashboard → Edge Functions que aparecen como "Active".',
    'Revisar los logs en Edge Functions → Logs para confirmar que ejecutan sin errores.',
    'Verificar en Database → pg_cron que los schedules están activos.',
]
for i, step in enumerate(ef_steps, 1):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(step)
    run.font.size = Pt(10)
    if 'supabase functions deploy' in step or 'supabase link' in step or 'supabase login' in step:
        run.font.name = 'Courier New'
        run.font.color.rgb = RGBColor(39, 174, 96)

doc.add_page_break()

# ── SECCIÓN 5 — CRON JOBS ─────────────────────────────────────────
add_heading(doc, '5. Implementaciones Faltantes — Cron Jobs', level=1)
add_colored_paragraph(doc,
    'Se definieron 11 cron jobs en vercel.json. Solo 5 tienen implementación en el código. '
    'Los 6 restantes devuelven 404 cuando Vercel intenta ejecutarlos.',
    size=11)
doc.add_paragraph()

# Tabla de cron jobs
add_heading(doc, '5.1. Estado actual de los Cron Jobs', level=2)
cron_rows = [
    ('sync-live-odds', '*/5 * * * *', 'Actualiza cuotas en tiempo real', '✓ Implementado', 'C8E6C9'),
    ('sync-fixtures', '0 */6 * * *', 'Sincroniza partidos cada 6h', '✓ Implementado', 'C8E6C9'),
    ('detect-value-bets', '*/10 * * * *', 'Detecta apuestas de valor', '✓ Implementado (Next.js fallback)', 'C8E6C9'),
    ('generate-parlays', '0 6,10,14 * * *', 'Genera parlays 3× al día', '✓ Implementado', 'C8E6C9'),
    ('expire-subscriptions', '0 5 * * *', 'Baja suscripciones expiradas', '✓ Implementado', 'C8E6C9'),
    ('generate-blog', '0 12 * * *', 'Genera contenido de blog con IA', '✗ Falta implementar', 'FFCDD2'),
    ('refresh-leaderboard', '0 * * * *', 'Refresca vista materializada', '✗ Falta implementar', 'FFCDD2'),
    ('sync-payu', '0 3 * * *', 'Reconcilia pagos PayU', '✗ Falta implementar', 'FFCDD2'),
    ('sync-results', '*/30 * * * *', 'Actualiza resultados de partidos', '✗ Falta implementar', 'FFCDD2'),
    ('cleanup-odds', '0 2 * * 0', 'Limpia cuotas antiguas (semanal)', '✗ Falta implementar', 'FFCDD2'),
    ('sync-injuries', '0 6,18 * * *', 'Sincroniza lesiones 2× al día', '✗ Falta implementar', 'FFCDD2'),
]

table = doc.add_table(rows=1, cols=4)
table.style = 'Table Grid'
hdr = table.rows[0].cells
for i, h in enumerate(['Cron Job', 'Schedule', 'Descripción', 'Estado']):
    shade_cell(hdr[i], '2C3E50')
    r = hdr[i].paragraphs[0].add_run(h)
    r.bold = True
    r.font.color.rgb = RGBColor(255, 255, 255)
    r.font.size = Pt(9)

for job, schedule, desc, estado, fill in cron_rows:
    row = table.add_row().cells
    row[0].text = job
    row[1].text = schedule
    row[2].text = desc
    row[3].text = estado
    shade_cell(row[3], fill)
    for j in range(4):
        for para in row[j].paragraphs:
            for run in para.runs:
                run.font.size = Pt(9)

doc.add_paragraph()
add_heading(doc, '5.2. Pasos para implementar los cron jobs faltantes', level=2)
add_colored_paragraph(doc,
    'Para cada cron job faltante, crear el archivo src/app/api/cron/[nombre]/route.ts con la siguiente estructura:',
    size=11)
doc.add_paragraph()

cron_missing_details = [
    ('refresh-leaderboard',
     'Llama al RPC refresh_leaderboard() de Supabase. Este RPC ya existe en la base de datos (migración 00003_functions.sql). Solo hace falta una ruta que lo invoque con el cliente admin.',
     'CRÍTICO — el leaderboard nunca se actualiza sin esto.'),
    ('sync-results',
     'Consulta la API de API-Football los resultados de partidos del día anterior y actualiza la columna result en la tabla matches. También llama a recalc_tipster_stats() para recalcular ROI de usuarios.',
     'CRÍTICO — los picks de usuarios nunca se resuelven sin esto.'),
    ('cleanup-odds',
     'Borra registros de odds con más de 30 días de antigüedad para mantener la base de datos ligera. Query: DELETE FROM odds WHERE created_at < NOW() - INTERVAL "30 days".',
     'IMPORTANTE — sin esto la tabla odds crece indefinidamente.'),
    ('sync-injuries',
     'Consulta el endpoint /injuries de API-Football para los próximos partidos y upserta en la tabla injuries. Usar con cuidado: consume cuota de la API (100 req/día en plan gratuito).',
     'MODERADO — mejora la calidad de predicciones.'),
    ('sync-payu',
     'Consulta el historial de transacciones de PayU para detectar pagos APPROVED que no llegaron por IPN. Crea las suscripciones faltantes. Requiere la API de reportes de PayU.',
     'IMPORTANTE — previene usuarios que pagaron pero no recibieron acceso.'),
    ('generate-blog',
     'Llama a la API de Claude/OpenAI para generar un artículo de blog basado en las noticias del día y los value bets detectados. Guarda en la tabla blog_posts.',
     'OPCIONAL para el MVP — puede postergarse.'),
]

for i, (nombre, descripcion, prioridad) in enumerate(cron_missing_details, 1):
    add_colored_paragraph(doc, f'{i}. {nombre}', size=11, bold=True)
    add_colored_paragraph(doc, f'   Qué hace: {descripcion}', size=10, color=(44, 62, 80))
    add_colored_paragraph(doc, f'   Prioridad: {prioridad}', size=10, color=(192, 57, 43), space_after=8)

doc.add_page_break()

# ── SECCIÓN 6 — PÁGINAS FALTANTES ────────────────────────────────
add_heading(doc, '6. Implementaciones Faltantes — Páginas y Rutas', level=1)
doc.add_paragraph()

pages_missing = [
    ('src/app/(app)/mis-picks/page.tsx',
     'Página donde el usuario registra y visualiza sus propias apuestas.',
     'Referenciada desde el dashboard. La tabla user_picks ya existe. El RPC register_pick() ya existe.',
     'CRÍTICO — el dashboard tiene un enlace que lleva a un 404.'),
    ('src/app/(app)/telegram/page.tsx',
     'Configuración del bot de Telegram para alertas Pro.',
     'El campo telegram_chat_id existe en profiles. Falta la UI para que el usuario vincule su chat ID.',
     'IMPORTANTE para features Pro.'),
    ('src/app/(app)/webhooks/page.tsx',
     'Gestión de webhooks personalizados para usuarios Pro.',
     'La tabla user_webhooks existe. Falta la UI para crear/eliminar/probar webhooks.',
     'IMPORTANTE para features Pro.'),
    ('src/app/api/track/affiliate/route.ts',
     'Endpoint que registra los clics en enlaces de afiliados.',
     'La tabla affiliate_clicks y los bookmakers están en la BD. La URL existe en vercel.json con rate limit configurado pero el archivo de ruta no existe.',
     'IMPORTANTE — sin esto no hay métricas de afiliados.'),
]

for i, (ruta, que_es, detalle, prioridad) in enumerate(pages_missing, 1):
    add_colored_paragraph(doc, f'{i}. {ruta}', size=10, bold=True, color=(0, 82, 163))
    add_colored_paragraph(doc, f'   Qué es: {que_es}', size=10)
    add_colored_paragraph(doc, f'   Detalle: {detalle}', size=10, color=(44, 62, 80))
    add_colored_paragraph(doc, f'   Prioridad: {prioridad}', size=10, color=(192, 57, 43), space_after=10)

doc.add_page_break()

# ── SECCIÓN 7 — LO QUE SÍ ESTÁ LISTO ────────────────────────────
add_heading(doc, '7. Lo que SÍ está listo', level=1)
add_colored_paragraph(doc,
    'Las siguientes áreas están completamente implementadas y no requieren cambios antes del lanzamiento:',
    size=11)
doc.add_paragraph()

ready_items = [
    ('Base de datos',
     '17 tablas con diseño maduro. RLS habilitado en todas las tablas. Triggers automáticos para timestamps, creación de perfiles y cálculo de probabilidades implícitas. 5 bookmakers y 10 ligas pre-cargados.'),
    ('Sistema de autenticación',
     'Middleware de Supabase con refresh automático de tokens. Separación correcta de cliente público vs admin. Funciones isPremiumUser() / isProUser() / getUserTier() funcionando.'),
    ('Seguridad',
     'Rate limiting en rutas sensibles. Detección y bloqueo de bots maliciosos. Verificación de firma en webhooks de Stripe y PayU. RLS impide acceso no autorizado a datos premium.'),
    ('Página de inicio (marketing)',
     'Muestra los próximos 48h de partidos. Value bets gratuitos + teaser de premium. Sección de características. Completamente funcional con datos reales de API-Football.'),
    ('Dashboard de usuario',
     'Estadísticas personales (ROI 30d, win rate, picks, favoritos). Parlays recientes. Sección de herramientas Pro visible solo para tier Pro. Mensajes de bienvenida post-checkout.'),
    ('Sistema de precios / premium gate',
     'Tres planes (Free, Premium, Pro) con features claramente diferenciadas. Botones duales COP (PayU) + USD (Stripe). La lógica de restricción de acceso funciona correctamente.'),
    ('Leaderboard',
     'Vista materializada con ranking de usuarios. Podio visual para top 3. Requiere mínimo 10 picks para aparecer. Resalta al usuario actual.'),
    ('Lógica matemática',
     'Conversión de cuotas a probabilidades implícitas (Multiplicative + método Shin). Detección de value bets (umbral ≥3% de ventaja). Modelo Poisson con datos xG. Kelly Criterion. Calculadora de parlays.'),
    ('Features Pro implementadas',
     'Backtesting: historial de apuestas resueltas con estadísticas de ROI. API Key: generación y rotación de clave personal. Página de precios con comparación detallada.'),
    ('Configuración de afiliados',
     '5 casas de apuestas configuradas con sus logos, colores y URLs de afiliado. La función buildAffiliateUrl() genera URLs correctas desde variables de entorno.'),
]

for titulo, descripcion in ready_items:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(5)
    run_bold = p.add_run(f'{titulo}: ')
    run_bold.bold = True
    run_bold.font.size = Pt(10)
    run_bold.font.color.rgb = RGBColor(39, 174, 96)
    run_normal = p.add_run(descripcion)
    run_normal.font.size = Pt(10)

doc.add_page_break()

# ── SECCIÓN 8 — PLAN DE ACCIÓN ────────────────────────────────────
add_heading(doc, '8. Plan de Acción por Semanas', level=1)
add_colored_paragraph(doc,
    'Este plan está ordenado por dependencias. Las tareas de la Semana 1 desbloquean las de la Semana 2, etc.',
    size=11)
doc.add_paragraph()

semanas = [
    ('Semana 1', 'Pagos y Entorno', 'CC0000', [
        ('Configurar Stripe', [
            'Crear cuenta o acceder al dashboard de Stripe',
            'Crear 4 productos: Premium Mensual, Premium Anual, Pro Mensual, Pro Anual',
            'Copiar los 4 Price IDs y llenarlos en .env.local',
            'Copiar Secret Key y Publishable Key',
            'Configurar webhook y copiar Signing Secret',
        ]),
        ('Generar CRON_SECRET', [
            'Abrir terminal y ejecutar: openssl rand -hex 32',
            'Copiar el resultado y pegarlo en .env.local como CRON_SECRET',
            'Agregar la misma variable en Vercel → Settings → Environment Variables',
        ]),
        ('Configurar PayU para producción', [
            'Registrar dominio en panel de PayU',
            'Llenar PAYU_RESPONSE_URL con la URL real del dominio',
            'Llenar PAYU_NOTIFICATION_URL con la URL del webhook',
            'Registrar el IPN en el panel de PayU',
        ]),
        ('Obtener claves de APIs externas', [
            'Registrarse en newsapi.org y copiar la API key',
            'Llenar NEWS_API_KEY en .env.local',
            'Decidir si se usa the-odds-api.com o se elimina esa dependencia',
        ]),
        ('Desplegar Supabase Edge Functions', [
            'Instalar CLI: npm install -g supabase',
            'Ejecutar: supabase login',
            'Ejecutar: supabase link --project-ref [PROJECT_ID]',
            'Ejecutar: supabase functions deploy detect-value-bets',
            'Ejecutar: supabase functions deploy generate-parlays',
            'Verificar en el dashboard de Supabase que están activas',
        ]),
    ]),
    ('Semana 2', 'Funcionalidades Faltantes Críticas', '0066CC', [
        ('Implementar página mis-picks', [
            'Crear src/app/(app)/mis-picks/page.tsx',
            'Listar apuestas del usuario desde tabla user_picks',
            'Mostrar estado (pendiente/ganado/perdido) y ROI',
            'Conectar con el RPC register_pick() para añadir nuevas apuestas',
        ]),
        ('Implementar rastreo de afiliados', [
            'Crear src/app/api/track/affiliate/route.ts',
            'Registrar cada clic en la tabla affiliate_clicks con: bookmaker_id, user_id (si está logueado), timestamp, IP anonimizada',
            'Retornar redirección 302 a la URL del afiliado',
        ]),
        ('Implementar cron refresh-leaderboard', [
            'Crear src/app/api/cron/refresh-leaderboard/route.ts',
            'Verificar el header Authorization con CRON_SECRET',
            'Llamar al RPC: SELECT refresh_leaderboard() con el cliente admin de Supabase',
            'Retornar 200 con el resultado',
        ]),
        ('Implementar cron sync-results', [
            'Crear src/app/api/cron/sync-results/route.ts',
            'Consultar /fixtures?status=FT de API-Football para partidos del día anterior',
            'Actualizar la columna result en la tabla matches',
            'Llamar a recalc_tipster_stats() para cada usuario afectado',
        ]),
        ('Implementar cron cleanup-odds', [
            'Crear src/app/api/cron/cleanup-odds/route.ts',
            'Eliminar registros de odds con más de 30 días: DELETE FROM odds WHERE created_at < NOW() - INTERVAL "30 days"',
            'Loggear cuántos registros se eliminaron',
        ]),
    ]),
    ('Semana 3', 'Testing y Verificación', '006633', [
        ('Test de pago Stripe end-to-end', [
            'Verificar que el botón "Pagar con USD" abre Stripe Checkout',
            'Completar pago con tarjeta de prueba: 4242 4242 4242 4242',
            'Verificar que el webhook recibe el evento checkout.session.completed',
            'Verificar que la tabla subscriptions se actualiza correctamente',
            'Verificar que el perfil del usuario cambia a tier "premium"',
            'Verificar que al acceder a contenido premium ya está desbloqueado',
        ]),
        ('Test de pago PayU end-to-end', [
            'Verificar que el botón "Pagar con COP" genera el formulario PayU',
            'Completar pago con tarjeta de prueba de PayU sandbox',
            'Verificar que el IPN llega a /api/webhooks/payu',
            'Verificar actualización de subscriptions y perfil',
        ]),
        ('Verificar seguridad RLS', [
            'Iniciar sesión como usuario gratuito',
            'Intentar acceder directamente a value bets premium via Supabase client',
            'Confirmar que la respuesta es una lista vacía (RLS bloqueando acceso)',
            'Verificar que usuario premium SÍ recibe los datos',
        ]),
        ('Verificar cron jobs en Vercel', [
            'Ir a Vercel Dashboard → proyecto → Functions → Cron',
            'Verificar que los 11 cron jobs aparecen configurados',
            'Revisar logs de las últimas ejecuciones',
            'Confirmar que no hay errores 404 o 500 en ningún job',
        ]),
        ('Ejecutar tests E2E con Playwright', [
            'Ejecutar: npm run test:e2e',
            'Verificar que todos los tests pasan',
            'Revisar el reporte de tests en playwright-report/',
        ]),
    ]),
    ('Semana 4', 'Lanzamiento', '4B0082', [
        ('Configuraciones finales de producción', [
            'Cambiar PAYU_ENV de "sandbox" a "production"',
            'Cambiar las claves de Stripe de test (sk_test_) a producción (sk_live_)',
            'Verificar que NEXT_PUBLIC_SITE_URL apunta al dominio real',
            'Verificar que todas las URLs de webhooks son las de producción',
        ]),
        ('Configurar analytics y monitoreo', [
            'Registrarse en posthog.com y llenar NEXT_PUBLIC_POSTHOG_KEY',
            'Configurar OneSignal y llenar APP_ID + REST API KEY (para notificaciones push)',
            'Configurar Sentry para monitoreo de errores (recomendado)',
        ]),
        ('Configurar Telegram bot (opcional MVP)', [
            'Crear bot con @BotFather en Telegram: /newbot',
            'Copiar el token y llenarlo como TELEGRAM_BOT_TOKEN',
            'Configurar el webhook del bot: POST https://api.telegram.org/bot[TOKEN]/setWebhook',
            'Completar la página src/app/(app)/telegram/page.tsx',
        ]),
        ('Verificación final antes de lanzar', [
            'Revisar el checklist completo de la Sección 10 de este documento',
            'Hacer una prueba de carga básica (100 usuarios simultáneos)',
            'Verificar que el sitio carga en menos de 3 segundos',
            'Confirmar que los cron jobs ejecutaron al menos una vez en producción',
            'Verificar que las Edge Functions de Supabase están activas y sin errores',
        ]),
        ('¡Lanzar!', [
            'Hacer el deploy final a Vercel: git push origin main',
            'Anunciar en redes sociales y canales de Telegram',
            'Monitorear los logs durante las primeras 24 horas',
            'Estar preparado para hotfixes inmediatos si algo falla',
        ]),
    ]),
]

for semana, titulo, color_hex, tareas in semanas:
    # Encabezado de semana
    r, g, b = int(color_hex[:2], 16), int(color_hex[2:4], 16), int(color_hex[4:], 16)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    run = p.add_run(f'  {semana}: {titulo}  ')
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(255, 255, 255)

    # Workaround: shade the paragraph background
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    pPr.append(shd)

    for tarea_titulo, pasos in tareas:
        add_colored_paragraph(doc, f'▶ {tarea_titulo}', size=11, bold=True, color=(r, g, b), space_before=6)
        for j, paso in enumerate(pasos, 1):
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.left_indent = Inches(0.3)
            p.paragraph_format.space_after = Pt(2)
            rr = p.add_run(paso)
            rr.font.size = Pt(10)

    doc.add_paragraph()

doc.add_page_break()

# ── SECCIÓN 9 — ESTIMACIÓN ────────────────────────────────────────
add_heading(doc, '9. Estimación de Tiempo', level=1)
doc.add_paragraph()

table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'
hdr = table.rows[0].cells
for i, h in enumerate(['Escenario', 'Duración estimada', 'Descripción']):
    shade_cell(hdr[i], '2C3E50')
    r = hdr[i].paragraphs[0].add_run(h)
    r.bold = True
    r.font.color.rgb = RGBColor(255, 255, 255)
    r.font.size = Pt(10)

est_rows = [
    ('MVP mínimo', '1–2 semanas',
     'Solo pagos (Stripe + PayU), Edge Functions desplegadas, crons críticos. Se posponen: Telegram, webhooks, blog automático, sync-injuries.'),
    ('Producto completo', '3–4 semanas',
     'Todos los crons, todas las páginas Pro, Telegram, blog automático, tests E2E completos, monitoreo.'),
    ('Lanzamiento suave (soft launch)', '2 semanas',
     'MVP funcional con pagos operativos, monitoreo básico y al menos 50% de features Pro implementadas.'),
]

for escenario, duracion, desc in est_rows:
    row = table.add_row().cells
    row[0].text = escenario
    row[1].text = duracion
    row[2].text = desc
    shade_cell(row[1], 'E3F2FD')
    for j in range(3):
        for para in row[j].paragraphs:
            for run in para.runs:
                run.font.size = Pt(10)

doc.add_paragraph()
add_colored_paragraph(doc,
    'Nota: Los tiempos asumen un desarrollador trabajando a tiempo completo. Con ayuda de Claude Code, '
    'las implementaciones de código pueden acelerarse significativamente.',
    size=10, color=(127, 140, 141))

doc.add_page_break()

# ── SECCIÓN 10 — CHECKLIST FINAL ─────────────────────────────────
add_heading(doc, '10. Checklist Final Pre-Lanzamiento', level=1)
add_colored_paragraph(doc,
    'Marcar cada item antes de hacer el deploy final a producción:',
    size=11)
doc.add_paragraph()

checklist_sections = [
    ('Variables de entorno', [
        'STRIPE_SECRET_KEY tiene valor real (sk_live_...)',
        'STRIPE_WEBHOOK_SECRET tiene valor real (whsec_...)',
        'STRIPE_PRICE_ID_MONTHLY tiene Price ID real de Stripe',
        'STRIPE_PRICE_ID_YEARLY tiene Price ID real de Stripe',
        'STRIPE_PRICE_ID_PRO_MONTHLY tiene Price ID real de Stripe',
        'STRIPE_PRICE_ID_PRO_YEARLY tiene Price ID real de Stripe',
        'CRON_SECRET tiene valor generado con openssl',
        'PAYU_RESPONSE_URL apunta al dominio real',
        'PAYU_NOTIFICATION_URL apunta al dominio real',
        'NEWS_API_KEY tiene clave real',
        'Todas las variables están también en Vercel → Environment Variables',
    ]),
    ('Pagos', [
        'Flujo Stripe funciona: botón → checkout → pago → tier upgrade',
        'Flujo PayU funciona: botón → formulario → pago → tier upgrade',
        'Webhook Stripe recibe eventos correctamente (verificar Stripe Dashboard → Webhooks)',
        'Webhook PayU (IPN) recibe eventos correctamente',
        'La suscripción se baja correctamente cuando el pago falla',
        'PAYU_ENV cambiado a "production" en el momento del lanzamiento',
    ]),
    ('Base de datos y Edge Functions', [
        'Migraciones aplicadas: supabase db push',
        'Edge Function detect-value-bets desplegada y activa',
        'Edge Function generate-parlays desplegada y activa',
        'Los pg_cron schedules están corriendo (verificar en Supabase → Database → pg_cron)',
        'La tabla leaderboard (vista materializada) tiene datos',
        'Los 5 bookmakers están seeded en la tabla bookmakers',
        'Las 10 ligas están seeded en la tabla leagues',
    ]),
    ('Cron Jobs (Vercel)', [
        'sync-live-odds ejecuta sin errores',
        'sync-fixtures ejecuta sin errores',
        'detect-value-bets ejecuta sin errores',
        'generate-parlays ejecuta sin errores',
        'expire-subscriptions ejecuta sin errores',
        'refresh-leaderboard implementado y ejecuta',
        'sync-results implementado y ejecuta',
        'cleanup-odds implementado y ejecuta',
    ]),
    ('Páginas y Funcionalidad', [
        'La página de inicio carga partidos reales de API-Football',
        'El registro de usuario crea perfil automáticamente',
        'El login funciona con email/password',
        'El dashboard muestra estadísticas del usuario',
        'La página /premium muestra los planes correctamente',
        'La página /mis-picks existe y carga sin errores',
        'El leaderboard muestra usuarios reales',
        'Los value bets premium están ocultos para usuarios gratuitos',
        'Los value bets premium son visibles para usuarios premium',
    ]),
    ('Seguridad', [
        'Usuario anónimo no puede leer value_bets vía Supabase client',
        'Usuario gratuito no puede leer parlays premium vía Supabase client',
        'Las rutas /dashboard, /premium/content redirigen a login si no hay sesión',
        'El rate limiting funciona (429 después de múltiples requests rápidos)',
        'Los webhooks rechazan requests sin firma válida',
        'La clave de servicio de Supabase NO está expuesta al cliente (verificar bundle)',
    ]),
    ('Performance', [
        'El sitio carga en menos de 3 segundos en 3G (verificar con PageSpeed Insights)',
        'Las imágenes están optimizadas (Next.js Image component usado)',
        'Las páginas estáticas tienen ISR configurado donde corresponde',
        'No hay N+1 queries en el dashboard (verificar logs de Supabase)',
    ]),
]

for section_title, items in checklist_sections:
    add_colored_paragraph(doc, section_title, size=12, bold=True, space_before=8)
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Inches(0.2)
        run = p.add_run(f'☐  {item}')
        run.font.size = Pt(10)

doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('— Documento generado automáticamente por auditoría de código —')
r.font.size = Pt(9)
r.font.color.rgb = RGBColor(149, 165, 166)
r.italic = True

# ── Guardar ───────────────────────────────────────────────────────
output_path = r'C:\Users\sunec\dev\apuestavalue\Plan_Lanzamiento_Coutazo.docx'
doc.save(output_path)
print(f'Documento guardado en: {output_path}')
