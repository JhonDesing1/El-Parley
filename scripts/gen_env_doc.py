from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# --- Estilos globales ---
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10)

# Márgenes
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

def heading(text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.runs[0]
    run.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)  # azul
    return p

def shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'

    # Cabecera
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        shade_cell(hdr_cells[i], '1E40AF')
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.bold = True
        run.font.size = Pt(9)

    # Filas
    for ri, row in enumerate(rows):
        cells = table.add_row().cells
        fill = 'EFF6FF' if ri % 2 == 0 else 'FFFFFF'
        for i, val in enumerate(row):
            cells[i].text = val
            shade_cell(cells[i], fill)
            p = cells[i].paragraphs[0]
            p.runs[0].font.size = Pt(9)

    # Anchos de columna
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)

    doc.add_paragraph()

# ── TÍTULO ──────────────────────────────────────────────────────────────────
title = doc.add_heading('ApuestaValue — Variables de Entorno', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.runs[0].font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)

doc.add_paragraph('Copia .env.example como .env.local y rellena los valores indicados a continuación.')
doc.add_paragraph()

# ── MÍNIMO PARA EJECUTAR ─────────────────────────────────────────────────────
heading('Mínimo indispensable para ejecutar localmente', level=2)
doc.add_paragraph(
    '1. Supabase (URL + anon key + service role key)\n'
    '2. API_FOOTBALL_KEY (datos de partidos)\n'
    '3. Stripe (keys de test + al menos un price ID)\n'
    '4. CRON_SECRET  →  generar con:  openssl rand -hex 32\n'
    '5. NEXT_PUBLIC_SITE_URL=http://localhost:3000  (ya está en el ejemplo)'
)
doc.add_paragraph()

# ── SITE ────────────────────────────────────────────────────────────────────
heading('Site / Next.js', level=2)
add_table(
    ['Variable', 'Valor por defecto / Descripción'],
    [
        ['NEXT_PUBLIC_SITE_URL', 'http://localhost:3000  (cambiar al dominio en producción)'],
        ['NEXT_PUBLIC_SITE_NAME', 'ApuestaValue'],
    ],
    col_widths=[7, 10]
)

# ── SUPABASE ─────────────────────────────────────────────────────────────────
heading('Supabase', level=2)
doc.add_paragraph('Dashboard → Project Settings → API')
add_table(
    ['Variable', 'Dónde obtenerla', 'Obligatoria'],
    [
        ['NEXT_PUBLIC_SUPABASE_URL', 'Dashboard → Project Settings → API', '✅'],
        ['NEXT_PUBLIC_SUPABASE_ANON_KEY', 'Dashboard → Project Settings → API', '✅'],
        ['SUPABASE_SERVICE_ROLE_KEY', 'Dashboard → Project Settings → API  ⚠️ solo server-side', '✅'],
        ['SUPABASE_PROJECT_ID', 'Dashboard → Project Settings → General', '✅'],
        ['SUPABASE_DB_PASSWORD', 'Contraseña que pusiste al crear el proyecto', '✅'],
    ],
    col_widths=[6, 9, 2.5]
)

# ── APIs DEPORTIVAS ──────────────────────────────────────────────────────────
heading('APIs de datos deportivos', level=2)
add_table(
    ['Variable', 'Dónde obtenerla', 'Obligatoria'],
    [
        ['API_FOOTBALL_KEY', 'api-football.com / RapidAPI — plan gratuito 100 req/día', '✅'],
        ['API_FOOTBALL_HOST', 'v3.football.api-sports.io  (valor fijo)', '✅'],
        ['ODDS_API_KEY', 'the-odds-api.com — 500 req/mes gratis', 'Opcional'],
        ['NEWS_API_KEY', 'newsapi.org — plan gratuito', 'Opcional'],
    ],
    col_widths=[6, 9, 2.5]
)

# ── STRIPE ───────────────────────────────────────────────────────────────────
heading('Stripe (pagos globales)', level=2)
doc.add_paragraph('Dashboard Stripe → Developers → API Keys / Webhooks / Products')
add_table(
    ['Variable', 'Dónde obtenerla', 'Obligatoria'],
    [
        ['STRIPE_SECRET_KEY', 'Dashboard → Developers → API Keys', '✅'],
        ['STRIPE_WEBHOOK_SECRET', 'Dashboard → Developers → Webhooks', '✅'],
        ['NEXT_PUBLIC_STRIPE_PUBLISHABLE_KEY', 'Dashboard → Developers → API Keys', '✅'],
        ['STRIPE_PRICE_ID_MONTHLY', 'Dashboard → Products (plan Premium mensual)', '✅'],
        ['STRIPE_PRICE_ID_YEARLY', 'Dashboard → Products (plan Premium anual)', '✅'],
        ['STRIPE_PRICE_ID_PRO_MONTHLY', 'Dashboard → Products (plan Pro mensual)', 'Opcional'],
        ['STRIPE_PRICE_ID_PRO_YEARLY', 'Dashboard → Products (plan Pro anual)', 'Opcional'],
    ],
    col_widths=[7, 8.5, 2]
)

# ── PAYU ─────────────────────────────────────────────────────────────────────
heading('PayU (Colombia / COP)', level=2)
doc.add_paragraph('Los valores de sandbox ya vienen en .env.example. Para producción, reemplaza con tus credenciales reales.')
add_table(
    ['Variable', 'Valor sandbox (pruebas)', 'Obligatoria'],
    [
        ['PAYU_ENV', 'sandbox', '✅'],
        ['PAYU_MERCHANT_ID', '508029', '✅'],
        ['PAYU_ACCOUNT_ID', '512321', '✅'],
        ['PAYU_API_KEY', '4Vj8eK4rloUd272L48hsrarnUA', '✅'],
        ['PAYU_API_LOGIN', 'pRRXKOl8ikMmt9u', '✅'],
        ['PAYU_RESPONSE_URL', 'https://tu-dominio.com/api/payu-response', '✅'],
        ['PAYU_NOTIFICATION_URL', 'https://tu-dominio.com/api/webhooks/payu', '✅'],
    ],
    col_widths=[6, 8.5, 3]
)

# ── NOTIFICACIONES ───────────────────────────────────────────────────────────
heading('Notificaciones', level=2)
add_table(
    ['Variable', 'Dónde obtenerla', 'Obligatoria'],
    [
        ['TELEGRAM_BOT_TOKEN', 'Crear bot con @BotFather en Telegram', 'Opcional'],
        ['NEXT_PUBLIC_TELEGRAM_BOT_USERNAME', 'Nombre elegido al crear el bot', 'Opcional'],
        ['NEXT_PUBLIC_ONESIGNAL_APP_ID', 'onesignal.com → Dashboard', 'Opcional'],
        ['ONESIGNAL_REST_API_KEY', 'onesignal.com → Settings → Keys & IDs', 'Opcional'],
    ],
    col_widths=[6.5, 8.5, 2.5]
)

# ── SEGURIDAD Y ANALYTICS ────────────────────────────────────────────────────
heading('Seguridad y Analytics', level=2)
add_table(
    ['Variable', 'Cómo obtenerla', 'Obligatoria'],
    [
        ['CRON_SECRET', 'Ejecutar en terminal:  openssl rand -hex 32', '✅'],
        ['NEXT_PUBLIC_POSTHOG_KEY', 'posthog.com → Project Settings', 'Opcional'],
        ['NEXT_PUBLIC_POSTHOG_HOST', 'https://app.posthog.com  (valor fijo)', 'Opcional'],
    ],
    col_widths=[6.5, 8.5, 2.5]
)

# ── AFILIADOS ────────────────────────────────────────────────────────────────
heading('Afiliados (puedes dejar los valores por defecto)', level=2)
add_table(
    ['Variable', 'Valor por defecto'],
    [
        ['NEXT_PUBLIC_BETPLAY_AFF', 'apuestavalue01'],
        ['NEXT_PUBLIC_WPLAY_AFF', 'apuestavalue02'],
        ['NEXT_PUBLIC_CODERE_AFF', 'apuestavalue03'],
        ['NEXT_PUBLIC_RIVALO_AFF', 'apuestavalue04'],
        ['NEXT_PUBLIC_1XBET_AFF', 'apuestavalue05'],
    ],
    col_widths=[7, 10.5]
)

doc.save('ApuestaValue_Variables_Entorno.docx')
print("Documento generado: ApuestaValue_Variables_Entorno.docx")
